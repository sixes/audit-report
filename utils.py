import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and PyInstaller."""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except AttributeError:
        # Use current directory in development
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def format_number(value, is_cost_or_admin=False, is_liability=False, is_tax=False):
    """Format a number: commas for thousands, parentheses for costs/admin/liabilities, dash for zero (others)."""
    if isinstance(value, str):
        value = int(value.replace(',', '').replace('(', '').replace(')', '')) if value else 0
    else:
        value = int(value)
    if value == 0:
            return "-"
    if is_tax:
        return f"{abs(value):,.0f}"
    elif is_cost_or_admin or is_liability:
        return f"({abs(value):,.0f})"
    else:
        if value < 0:
            return f"({abs(value):,.0f})"
        else:
            return f"{value:,.0f}"

def update_fields(doc):
    """Update all fields in the document, including page numbers in footers."""
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            for run in footer.runs:
                if 'PAGE' in run.text.upper():
                    run.text = ''
                    field = OxmlElement('w:fldSimple')
                    field.set(qn('w:instr'), 'PAGE')
                    run._element.append(field)

def insert_page_break_before_income_statement(doc):
    """Insert a page break before the Statement of Comprehensive Income table."""
    for i, para in enumerate(doc.paragraphs):
        if "Statement of Comprehensive Income" in para.text:
            prev_para = doc.paragraphs[i-1] if i > 0 else doc.add_paragraph()
            run = prev_para.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._element.append(br)
            break