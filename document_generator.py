import logging
import os
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docxtpl import DocxTemplate
from data_loader import DataLoader
from utils import resource_path, format_number, update_fields, insert_page_break_before_income_statement
from exceptions import InvalidTBSheetFormatError, UnrecognizedItemError, InvalidItemNameError, NetAssetsEquityMismatchError

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_generator.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class DocumentGenerator:
    FILE_TPLS = {
        "LAI_1":  "template/temp_first_lai.docx",
        "LEUNG_1": "template/temp_first_leung.docx",
        "WOCP_1": "template/temp_first_wocp.docx",
        "WH_1": "template/temp_first_wh.docx",
        "WU_1":   "template/temp_first_wu.docx",
        "WH":   "template/temp_not_first_wh.docx",
        "WOCP": "template/temp_not_first_wocp.docx",
        "LAI":  "template/temp_not_first_lai.docx",
        "WU":   "template/temp_not_first_wu.docx",
        "LEUNG": "template/temp_not_first_leung.docx",
    }
    AUX_TPLS = {
        "WH":   "template/temp_aux_wh.docx",
        "WOCP": "template/temp_aux_wocp.docx",
        "LAI":  "template/temp_aux_lai.docx",
        "WU":   "template/temp_aux_wu.docx",
        "LEUNG": "template/temp_aux_leung.docx",
    }

    due_from_final_holding_parent_company_items = [
        'amount due from final holding parent company',
        'amount due from a final holding parent company',
        'amount due from the final holding parent company',
        'amount due from final holding parent companies',
    ]
    due_to_final_holding_parent_company_items = [
        'amount due to final holding parent company',
        'amount due to a final holding parent company',
        'amount due to the final holding parent company',
        'amount due to final holding parent companies',
    ]
    due_final_holding_parent_company_items = [] + due_from_final_holding_parent_company_items + due_to_final_holding_parent_company_items

    due_from_shareholder_items = [
        'amount due from a shareholder', 
        'amount due from the shareholder', 
        'amount due from shareholder', 
        'amount due from shareholders', 
    ]
    due_to_shareholder_items = [
        'amount due to a shareholder', 
        'amount due to the shareholder', 
        'amount due to shareholder', 
        'amount due to shareholders'
    ]
    due_shareholder_items = [] + due_from_shareholder_items + due_to_shareholder_items

    due_from_imme_parent_company_items = [
        'amount due from an immediate parent company', 
        'amount due from the immediate parent company', 
        'amount due from immediate parent company', 
        'amount due from immediate parent companies'
    ]
    due_to_imme_parent_company_items = [
        'amount due to an immediate parent company', 
        'amount due to the immediate parent company', 
        'amount due to immediate parent company', 
        'amount due to immediate parent companies'
    ]
    due_imme_parent_company_items = [] + due_from_imme_parent_company_items + due_to_imme_parent_company_items

    due_from_ultimate_holding_company_items = [
        'amount due from an ultimate holding company', 
        'amount due from the ultimate holding company', 
        'amount due from ultimate holding company', 
        'amount due from ultimate holding companies',
    ] 
    due_to_ultimate_holding_company_items = [
        'amount due to an ultimate holding company', 
        'amount due to the ultimate holding company', 
        'amount due to ultimate holding company', 
        'amount due to ultimate holding companies'
    ]
    due_ultimate_holding_company_items = [] + due_from_ultimate_holding_company_items + due_to_ultimate_holding_company_items

    due_from_holding_company_items = [
        'amount due from a holding company', 
        'amount due from the holding company', 
        'amount due from holding company', 
        'amount due from holding companies', 
    ]
    due_to_holding_company_items = [
        'amount due to a holding company', 
        'amount due to the holding company', 
        'amount due to holding company', 
        'amount due to holding companies',
    ]
    due_holding_company_items = [] + due_from_holding_company_items + due_to_holding_company_items

    def __init__(self, category_manager):
        self._last_day_date = None
        self._last_day_date_num = None
        self._last_day_date_cn = None
        self._one_year_ago = None
        self._first_year = False
        self._cu_year_for_first_year = ""
        self._c_year_for_first_year = ""
        self._pr_year_for_first_year = ""
        self._over18m = False
        self._exactly_one_year_ago = None
        self._one_year_ago_cn = None
        self._date_of_incorporation=None
        self._audit_year = None
        self._is_december = None
        self._directors_list = None
        self._shareholders_list = None
        self._business_type = None
        self._company_name_en = None
        self._currency = None
        self._has_inventories_curr = False
        self._due_to_directors_curr = 0
        self._has_due_from_directors_curr = False
        self._due_from_directors_curr = 0
        self._inventories_curr = 0
        self._inventories_prev = 0
        self._closing_inventories_prev = 0
        self._closing_inventories_curr = 0
        self._has_subsidiary = False
        self._all_items = None
        self._all_items_curr = None
        self._all_items_prev = None
        self._accountant_helper = None
        self._statement_current = None
        self._statement_previous = None
        self._balance_current = None
        self._balance_previous = None
        self._category_manager = category_manager

    def get_due_info(self, due_from_items, due_to_items, due_all_items, need_title=True):
        # Original calculation for due_final_curr, due_final_prev
        due_final_curr = 0
        due_final_prev = 0

        # Current year: Check assets for "from" items
        for item in self._balance_current['current_assets']:
            if item['name'] in due_from_items:
                due_final_curr = item['value']
                break

        # Current year: Check liabilities for "to" items
        for item in self._balance_current['current_liabilities']:
            if item['name'] in due_to_items:
                due_final_curr = -item['value']
                break

        # Previous year: Check assets for "from" items
        for item in self._balance_previous['current_assets']:
            if item['name'] in due_from_items:
                due_final_prev = item['value']
                break

        # Previous year: Check liabilities for "to" items
        for item in self._balance_previous['current_liabilities']:
            if item['name'] in due_to_items:
                due_final_prev = -item['value']
                break

        # Calculate max
        due_final_max = max(due_final_curr, due_final_prev)

        # Format values
        due_final_curr_formatted = format_number(due_final_curr, is_cost_or_admin=False,
                                                 is_liability=False) if due_final_curr != 0 else "-"
        due_final_prev_formatted = format_number(due_final_prev, is_cost_or_admin=False,
                                                 is_liability=False) if due_final_prev != 0 else "-"
        due_final_max_formatted = format_number(due_final_max, is_cost_or_admin=False,
                                                is_liability=False) if due_final_max != 0 else "-"
        if due_final_curr == 0 and due_final_prev == 0:
            return {
                'need_footnote': False,
                'both_to': False,
                'curr': due_final_curr_formatted,
                'prev': due_final_prev_formatted,
                'max': due_final_max_formatted,
                'title_name': ''
            }

        title_name = ""
        if need_title:
            matches = [item for item in due_all_items if item in self._all_items]
            from_name = next((item for item in matches if item in due_from_items), "")
            to_name = next((item for item in matches if item in due_to_items), "")
            if from_name and to_name:
                title_name = from_name.replace('from', 'from/(to)')
            else:
                title_name = from_name if from_name else to_name

        if True: #self._first_year:
            return {
                'need_footnote': any(item in self._all_items_curr for item in due_all_items),
                'both_to': any(item in self._all_items_curr for item in due_all_items),
                'curr': due_final_curr_formatted,
                'prev': due_final_prev_formatted,
                'max': due_final_max_formatted,
                'title_name': title_name
            }

        # Logic for both_to
        curr_to = any(item in self._all_items_curr for item in due_to_items)
        prev_to = any(item in self._all_items_prev for item in due_to_items)
        both_to = (curr_to and prev_to) or (curr_to and due_final_prev == 0)

        # Logic for need_footnote
        # False if: (1) due_final_curr is 0 and prev has due_to items, or
        #           (2) no due_to/from items in both years
        need_footnote = not (due_final_curr == 0 and prev_to)

        return {
            'need_footnote': need_footnote,
            'both_to': both_to,
            'curr': due_final_curr_formatted,
            'prev': due_final_prev_formatted,
            'max': due_final_max_formatted,
            'title_name': title_name
        }

    def _initialize_common_data(
        self,
        last_day_of_year,
        business_type,
        company_name_en,
        directors,
        shareholders,
        currency,
        excel_file=None,
        first_year=False,
        current_year=None,
        date_of_incorporation=None
    ):
        """Initialize shared data used by both aux and main document generation."""
        # Validate directors
        if not directors or not isinstance(directors, list):
            logger.error("Directors list is empty or invalid")
            raise ValueError("Directors list cannot be empty or invalid")
        self._directors_list = [d.strip() for d in directors if d.strip()]
        if not self._directors_list:
            logger.error("Directors list is empty after cleaning")
            raise ValueError("Directors list cannot be empty after cleaning")

        self._shareholders_list = shareholders
        # Store other parameters
        self._business_type = business_type.lower()
        self._company_name_en = company_name_en
        self._currency = currency
        self._first_year = first_year
        self._date_of_incorporation = date_of_incorporation

        # Date calculations
        try:
            self._last_day_date = datetime.strptime(last_day_of_year, "%d %B %Y")
            #self._exactly_one_year_ago = self._last_day_date.replace(year=self._last_day_date.year - 1).strftime("%d %B %Y").lstrip("0")
            self._exactly_one_year_ago = self._last_day_date.replace(year=self._last_day_date.year - 1).strftime(
                "%d %B %Y")
            self._last_day_date_num = self._last_day_date.strftime("%d.%m.%Y").replace(" 0", "").replace(".0", ".")
            self._last_day_date_cn = self._last_day_date.strftime("%Y年%m月%d日")
            self._is_december = self._last_day_date.month == 12
            if self._is_december:
                # For December dates, set one_year_ago to 1 January of the same year
                one_year_ago_date = datetime(self._last_day_date.year, 1, 1)
            else:
                # For non-December dates, get next day and subtract one year
                next_day_date = self._last_day_date + timedelta(days=1)
                try:
                    one_year_ago_date = next_day_date.replace(year=next_day_date.year - 1)
                except ValueError:
                    # Handle leap year edge case (e.g., 29 Feb on non-leap year)
                    one_year_ago_date = next_day_date.replace(year=next_day_date.year - 1, day=next_day_date.day - 1)
            self._one_year_ago = one_year_ago_date.strftime("%d %B %Y").lstrip("0")
            self._one_year_ago_cn = one_year_ago_date.strftime("%Y年%m月%d日")

            if self._last_day_date.month >= 4:
                self._audit_year = f"{self._last_day_date.year}/{str(self._last_day_date.year + 1)[-2:]}"
            else:
                self._audit_year = f"{self._last_day_date.year - 1}/{str(self._last_day_date.year)[-2:]}"

            if self._first_year and self._date_of_incorporation:
                self._cu_year_for_first_year = self._last_day_date.strftime("%d/%m/%Y").lstrip("0").replace("/0", "/")
                incorporation_date = datetime.strptime(self._date_of_incorporation, "%d %B %Y")
                self._pr_year_for_first_year = incorporation_date.strftime("%d/%m/%Y").lstrip("0").replace("/0", "/")
                self._one_year_ago_cn = incorporation_date.strftime("%Y年%m月%d日")
                eighteen_months_later = incorporation_date + relativedelta(months=18)
                self._over18m = self._last_day_date >= eighteen_months_later
                self._c_year_for_first_year = self._last_day_date.year
        # Handle invalid date format    
        except ValueError as e:
            logger.error(f"Invalid date format for LastDayOfYear: {last_day_of_year}. Expected format: '31 December 2024'")
            raise ValueError(f"Invalid date format for LastDayOfYear: {last_day_of_year}. Expected format: '31 December 2024'")

        # Initialize trial balance data if provided
        if excel_file and current_year:
            tb_file = excel_file if excel_file else "example_tb_for_test.xlsx"
            self._accountant_helper = DataLoader(
                excel_file=tb_file,
                first_year=first_year,
                current_year=current_year,
                non_current_assets= self._category_manager.categories['non_current_assets'],
                current_assets= self._category_manager.categories['current_assets'],
                current_liabilities= self._category_manager.categories['current_liabilities'],
                non_current_liabilities= self._category_manager.categories['non_current_liabilities'],
                equity= self._category_manager.categories['equity'],
                revenue_items= self._category_manager.categories['revenue_items'],
                cost_of_sales_items= self._category_manager.categories['cost_of_sales_items'],
                closing_inventories= self._category_manager.categories['closing_inventories'],
                other_income_items= self._category_manager.categories['other_income_items'],
                general_admin_expenses_items= self._category_manager.categories['general_admin_expenses_items'],
                finance_costs_items= self._category_manager.categories['finance_costs_items'],
                tax_items= self._category_manager.categories['tax_items']
            )
            self._statement_current = self._accountant_helper.get_income_statement(current_year)
            self._balance_current = self._statement_current['BalanceSheet']
            self._statement_previous = self._accountant_helper.get_income_statement(current_year - 1)
            self._balance_previous = self._statement_previous['BalanceSheet']
            self._all_items_curr = set(str(row['Item']).strip().lower() for _, row in self._accountant_helper.data[current_year].iterrows())
            self._all_items_prev = set(str(row['Item']).strip().lower() for _, row in self._accountant_helper.data[current_year - 1].iterrows())

            # Calculate HasInventoriesCurr and InventoriesCurr
            self._has_inventories_curr = 'inventories' in self._all_items_curr
            self._inventories_curr = 0
            self._inventories_prev = 0
            for item in self._balance_current['current_assets']:
                if item['name'] == "inventories":
                    self._inventories_curr = item['value']
                    break
            for item in self._balance_previous['current_assets']:
                if item['name'] == "inventories":
                    self._inventories_prev = item['value']
                    break

            # Calculate HasDueFromDirectorsCurr and DueFromDirectorsCurr
            due_from_directors_items = ['amount due from director', 'amount due from a director', 'amount due from the director', 'amount due from directors']
            self._has_due_from_directors_curr = any(item in self._all_items_curr for item in due_from_directors_items)
            self._due_from_directors_curr = 0
            for item in self._balance_current['current_assets']:
                if item['name'] in due_from_directors_items:
                    self._due_from_directors_curr = item['value']
                    break

            # Calculate DueToDirectorsCurr
            self._due_to_directors_curr = 0
            due_to_directors_items = ['amount due to director', 'amount due to a director', 'amount due to the director', 'amount due to directors']
            for item in self._balance_current['current_liabilities']:
                if item['name'] in due_to_directors_items:
                    self._due_to_directors_curr = -item['value']
                    break

            # Calculate HasSubsidiary
            subsidiary_items = [
                'investments in a subsidiary', 'investments in subsidiaries',
                'interests in subsidiaries', 'interests in a subsidiary'
            ]
            self._has_subsidiary = any(item in self._all_items_curr for item in subsidiary_items)

    def generate_aux_document(
        self,
        last_day_of_year,
        business_type,
        aux_output_path,
        company_name_en,
        directors,
        shareholders,
        currency,
        has_stocking_letter,
        br_no,
        excel_file,
        current_year=None,
        first_year=False,
        company_name_cn=None,
        company_address="123 Business Road, Hong Kong",
        business_description="Technology and Software Development",
        additional_business_description="",
        audit_firm="Deloitte",
        approval_date="1st April 2035",
        auditor_name="Auditor Name",
        auditor_license="CPA12345",
        currency_desc="Hong Kong dollars",
        currency_full_desc="Hong Kong dollars (HK$)",
        shares_curr="10000",
        shares_prev="10000",
        has_name_changed=False,
        passed_date="",
        new_company_name="",
        effective_date="",
        old_company_name="",
        has_related_party=False,
        inventory_valuation="FIFO",
        tax_opt="1",
        capital_increase=True,
        has_ultimate_company=False,
        ultimate_company_option="",
        ultimate_company_name1="",
        ultimate_company_location1="",
        ultimate_company_name2="",
        ultimate_company_location2="",
        investment_in_company=False,
        investment_in_security=False,
        audit_opinion="Opinion",
        audit_type="",
        date_of_incorporation=None
    ):
        template_path = DocumentGenerator.AUX_TPLS[audit_type]
        if not os.path.exists(resource_path(template_path)):
            logger.error(f"Aux template file not found at: {template_path}")
            #raise FileNotFoundError(f"Aux template file not found at: {template_path}")
            return None, f"Aux template file not found at: {template_path}"

        logger.info(f"Attempting to load aux template: {template_path}")
        try:
            template = DocxTemplate(resource_path(template_path))
        except Exception as e:
            logger.error(f"Failed to initialize aux DocxTemplate: {str(e)}")
            raise ValueError(f"Failed to initialize aux DocxTemplate: {str(e)}")

        # Initialize common data if not already set
        self._initialize_common_data(
            last_day_of_year,
            business_type,
            company_name_en,
            directors,
            shareholders,
            currency,
            excel_file,
            first_year,
            current_year,
            date_of_incorporation
        )

        company_address_cleaned = company_address.replace('\n', ' ').strip()

        data = {
            "LastDayOfYear": last_day_of_year,
            "LastDayOfYearNum": self._last_day_date_num,
            "LastDayOfYearCN": self._last_day_date_cn,
            "OneYearAgoCN": self._one_year_ago_cn,
            "AuditYear": self._audit_year,
            "BusinessType": self._business_type,
            "CompanyNameInEnglishPlaceholder": self._company_name_en,
            "CompanyNameInChinesePlaceholder": company_name_cn or "",
            "CompanyAddressPlaceHolder": company_address_cleaned,
            "BusinessDescriptionPlaceholder": business_description,
            "bizAdditionalDesc": additional_business_description,
            "Directors": self._directors_list,
            "Shareholders": self._shareholders_list,
            "Currency": self._currency,
            "CurrencyDesc": currency_desc,
            "CurrencyFullDesc": currency_full_desc,
            "HasStockingLetter": has_stocking_letter,
            "BRNo": br_no,
            "HasInventoriesCurr": self._has_inventories_curr,
            "HasDueToDirectorsCurr": self._due_to_directors_curr != 0,
            "DueToDirectorsCurr": format_number(self._due_to_directors_curr, is_cost_or_admin=False, is_liability=False) if self._due_to_directors_curr != 0 else "-",
            "DueToDirectorsCurrFn": format_number(self._due_to_directors_curr, is_cost_or_admin=False, is_liability=False, is_tax=True) if self._due_to_directors_curr != 0 else "-",
            "HasDueFromDirectorsCurr": self._has_due_from_directors_curr,
            "DueFromDirectorsCurr": format_number(self._due_from_directors_curr, is_cost_or_admin=False, is_liability=False) if self._due_from_directors_curr != 0 else "-",
            "DueFromDirectorsCurrFn": format_number(self._due_from_directors_curr, is_cost_or_admin=False, is_liability=False, is_tax=True) if self._due_from_directors_curr != 0 else "-",
            "InventoriesCurr": format_number(self._inventories_curr, is_cost_or_admin=False, is_liability=False) if self._inventories_curr != 0 else "-",
            "HasSubsidiary": self._has_subsidiary,
            "AuditFirmInEnglishPlacehoder": audit_firm,
            "ApprovalDatePlaceholder": approval_date,
            "AuditorNamePlaceholder": auditor_name,
            "AuditorLicenseNoPlaceholder": auditor_license,
            "SharesCurr": format_number(shares_curr, is_cost_or_admin=False, is_liability=False),
            "SharesPrev": format_number(shares_prev, is_cost_or_admin=False, is_liability=False),
            "HasNameChanged": has_name_changed,
            "PassedDate": passed_date,
            "NewCompanyName": new_company_name,
            "EffectiveDate": effective_date,
            "OldCompanyName": old_company_name,
            "HasRelatedParty": has_related_party,
            "InventoryValuation": inventory_valuation,
            "TaxOpt": tax_opt,
            "CapitalIncrease": capital_increase,
            "HasUltimateCompany": has_ultimate_company,
            "UltimateCompanyOption": ultimate_company_option,
            "UltimateCompanyName1": ultimate_company_name1,
            "UltimateCompanyLocation1": ultimate_company_location1,
            "UltimateCompanyName2": ultimate_company_name2,
            "UltimateCompanyLocation2": ultimate_company_location2,
            "InvestmentInCompany": investment_in_company,
            "InvestmentInSecurity": investment_in_security,
            "Opinion": audit_opinion
        }

        excluded_fields = ["bizAdditionalDesc", "PassedDate", "NewCompanyName", "OldCompanyName", "EffectiveDate",
        "UltimateCompanyName2", "UltimateCompanyLocation2", "UltimateCompanyName1", "UltimateCompanyLocation1", 
        "UltimateCompanyOption", "InvestmentInCompany", "InvestmentInSecurity", "Opinion", "CompanyNameInChinesePlaceholder"]
        for key, value in data.items():
            if key not in excluded_fields and isinstance(value, str) and not value.strip():
                logger.error(f"Please fill in all fields: {key}")
                return None, f"Error: Please fill in all fields: {key}"
        # Render and save the template
        try:
            logger.info(f"Rendering aux template to {aux_output_path}")
            template.render(data)
            template.save(aux_output_path)
            logger.info(f"Aux document successfully generated at: {aux_output_path}")
            return True, ""
        except Exception as e:
            logger.error(f"Failed to render or save aux document: {str(e)}")
            return False, f"Error: Failed to generate aux document: {str(e)}"

    def generate_document(
        self,
        business_type="general trading",
        excel_file=None,
        output_path=None,
        current_year=None,
        first_year=False,
        company_name_en="Trump Company",
        company_name_cn=None,
        company_address="123 Business Road, Hong Kong",
        business_description="Technology and Software Development",
        additional_business_description="",
        last_day_of_year=None,
        date_of_incorporation=None,
        audit_firm="Deloitte",
        approval_date="1st April 2035",
        auditor_name="Auditor Name",
        auditor_license="CPA12345",
        currency="HK$",
        currency_desc="Hong Kong dollars",
        currency_full_desc="Hong Kong dollars (HK$)",
        directors=None,
        shares_curr="10000",
        shares_prev="10000",
        non_current_assets=None,
        current_assets=None,
        current_liabilities=None,
        non_current_liabilities=None,
        equity=None,
        revenue_items=None,
        cost_of_sales_items=None,
        closing_inventories=None,
        other_income_items=None,
        general_admin_expenses_items=None,
        finance_costs_items=None,
        tax_items=None,
        has_name_changed=False,
        passed_date="",
        new_company_name="",
        effective_date="",
        old_company_name="",
        has_related_party=False,
        inventory_valuation="FIFO",
        tax_opt="1",
        capital_increase="",
        has_ultimate_company=False,
        ultimate_company_option="",
        ultimate_company_name1="",
        ultimate_company_location1="",
        ultimate_company_name2="",
        ultimate_company_location2="",
        investment_in_company=False,
        investment_in_security=False,
        audit_opinion="Opinion",
        audit_type="WH",
        shareholders=None
    ):
        try:
            file_key = audit_type
            if first_year:
                file_key = audit_type + "_1"
            template_path = DocumentGenerator.FILE_TPLS[file_key]
            if not os.path.exists(resource_path(template_path)):
                logger.error(f"Template file not found at: {template_path}")
                raise FileNotFoundError(f"Template file not found at: {template_path}")

            logger.info(f"Attempting to load template: {template_path}")
            template = DocxTemplate(resource_path(template_path))
            if template is None:
                logger.error("Failed to initialize DocxTemplate: template is None")
                raise ValueError("Failed to initialize DocxTemplate: template is None")

            # Initialize common data if not already set
            self._initialize_common_data(
                last_day_of_year,
                business_type,
                company_name_en,
                directors,
                shareholders,
                currency,
                excel_file,
                first_year,
                current_year,
                date_of_incorporation
            )

            first_director_name = self._directors_list[0] if self._directors_list else ""

            if current_year is None:
                logger.error("current_year must be provided to generate_document")
                raise ValueError("current_year must be provided to generate_document")

            company_address_cleaned = company_address.replace('\n', ' ').strip()

            if self._business_type == "general trading":
                revenue_name = "Sales of goods"
            elif self._business_type == "services":
                revenue_name = "Services fee income"
            elif self._business_type == "agency services":
                revenue_name = "Agency service income"
            elif self._business_type in ("dormant", "investment holding"):
                revenue_name = "Sales of goods"
            else:
                logger.error("Business type must be 'general trading', 'services', 'dormant', 'agency services', or 'investment holding'")
                raise ValueError("Business type must be 'general trading', 'services', 'dormant', 'agency services', or 'investment holding'")

            # Use existing accountant_helper if initialized
            if not self._accountant_helper:
                tb_file = excel_file if excel_file else "example_tb_for_test.xlsx"
                logger.info(f"Using trial balance file: {tb_file}")
                self._accountant_helper = DataLoader(
                    excel_file=tb_file,
                    current_year=current_year,
                    non_current_assets=self._category_manager.categories['non_current_assets'],
                    current_assets=self._category_manager.categories['current_assets'],
                    current_liabilities=self._category_manager.categories['current_liabilities'],
                    non_current_liabilities=self._category_manager.categories['non_current_liabilities'],
                    equity=self._category_manager.categories['equity'],
                    revenue_items=self._category_manager.categories['revenue_items'],
                    cost_of_sales_items=self._category_manager.categories['cost_of_sales_items'],
                    closing_inventories=self._category_manager.categories['closing_inventories'],
                    other_income_items=self._category_manager.categories['other_income_items'],
                    general_admin_expenses_items=self._category_manager.categories['general_admin_expenses_items'],
                    finance_costs_items=self._category_manager.categories['finance_costs_items'],
                    tax_items=self._category_manager.categories['tax_items']
                )

            statement_current = self._accountant_helper.get_income_statement(current_year)
            previous_year = current_year - 1
            statement_previous = self._accountant_helper.get_income_statement(previous_year)

            balance_current = statement_current['BalanceSheet']
            balance_previous = statement_previous['BalanceSheet']

            all_items = set()
            for year in [current_year, previous_year]:
                df = self._accountant_helper.data[year]
                all_items.update(str(row['Item']).strip().lower() for _, row in df.iterrows())
            self._all_items = all_items

            logger.debug(f"all_items: {sorted(all_items)}")

            due_from_directos_items = ['amount due from director', 'amount due from a director', 'amount due from the director', 'amount due from directors']
            due_from_dir = any(item in all_items for item in due_from_directos_items)
            
            has_service_fee_income = 'services fee income' in all_items
            has_agency_fee_income = 'agency fee income' in all_items
            has_sales_of_goods = 'sales of goods' in all_items
            has_long_term_investments = 'long-term investments' in all_items
            has_property = 'property, plant and equipment' in all_items
            has_investment = any(item in all_items for item in [
                'investments in an associate', 'long-term investments', 'current investments'
            ])
            has_inventories = 'inventories' in all_items
            has_intangible_asset = 'intangible assets' in all_items
            has_reserve = any(item in all_items for item in ['capital reserves', 'reserves'])
            has_current_investments = 'current investments' in all_items
            has_intangible_assets = 'intangible assets' in all_items
            has_associate = any(item in all_items for item in ['investments in an associate', 'investments in associates'])

            logger.debug(f"has_inventories: {has_inventories}")
            logger.debug(f"due_from_dir: {due_from_dir}")
            logger.debug(f"has_associate: {has_associate}")
            logger.debug(f"has_long_term_investments: {has_long_term_investments}")
            logger.debug(f"has_property: {has_property}")
            logger.debug(f"has_investment: {has_investment}")
            logger.debug(f"has_intangible_asset: {has_intangible_asset}")

            pbt_current = statement_current['ProfitBeforeTax']
            pbt_previous = statement_previous['ProfitBeforeTax']
            if self._first_year:
               if pbt_current >=0:
                   plbft_name = "Profit"
               else:
                   plbft_name = "Loss"
            else:
                if pbt_current >= 0 and pbt_previous >= 0:
                    plbft_name = "Profit"
                elif pbt_current < 0 and pbt_previous < 0:
                    plbft_name = "Loss"
                else:
                    plbft_name = "Profit/(Loss)"

            pfy_current = statement_current['ProfitForYear']
            pfy_previous = statement_previous['ProfitForYear']
            if self._first_year:
                if pfy_current >= 0:
                    pl_name = "Profit"
                else:
                    pl_name = "Loss"
            else:
                if pfy_current >= 0 and pfy_previous >= 0:
                    pl_name = "Profit"
                elif pfy_current < 0 and pfy_previous < 0:
                    pl_name = "Loss"
                else:
                    pl_name = "Profit/(Loss)"

            pl_name_curr = "Profit" if pfy_current >= 0 else "Loss"

            pl_name_cap = ''.join(c.upper() if c.isalpha() else c for c in pl_name)

            gross_profit_current = statement_current['GrossProfit']
            gross_profit_previous = statement_previous['GrossProfit']
            if gross_profit_current >= 0 and gross_profit_previous >= 0:
                gross_pl_name = "PROFIT"
            elif gross_profit_current < 0 and gross_profit_previous < 0:
                gross_pl_name = "LOSS"
            else:
                gross_pl_name = "PROFIT/(LOSS)"

            net_assets_current = balance_current['net_assets']
            net_assets_previous = balance_previous['net_assets']
            if self._first_year:
                if net_assets_current >= 0:
                    net_assets_name = "Net assets"
                else:
                    net_assets_name = "Net liabilities"
            else:
                if net_assets_current > 0 and net_assets_previous > 0:
                    net_assets_name = "Net assets"
                elif net_assets_current < 0 and net_assets_previous < 0:
                    net_assets_name = "Net liabilities"
                else:
                    net_assets_name = "Net assets/(liabilities)"

            cash_bank = "-"
            for item in balance_current['current_assets']:
                if item['name'] == "cash and bank balances":
                    cash_bank = format_number(item['value'])
                    break
            
            long_term_investments_curr = 0
            long_term_investments_prev = 0
            if has_long_term_investments:
                for item in balance_current['non_current_assets']:
                    if item['name'] == "long-term investments":
                        long_term_investments_curr = item['value']
                        break
                for item in balance_previous['non_current_assets']:
                    if item['name'] == "long-term investments":
                        long_term_investments_prev = item['value']
                        break

            current_investment_curr = 0
            current_investment_prev = 0
            if has_current_investments:
                for item in balance_current['current_assets']:
                    if item['name'] == "current investments":
                        current_investment_curr = item['value']
                        break
                for item in balance_previous['current_assets']:
                    if item['name'] == "current investments":
                        current_investment_prev = item['value']
                        break

            audit_fee_current = "-"
            audit_fee_previous = "-"
            for item in statement_current['GeneralAdminExpensesDetails']:
                if item['name'] in ["audit fee", "auditors' remuneration"]:
                    audit_fee_current = format_number(item['value'], is_cost_or_admin=False)
                    break
            for item in statement_previous['GeneralAdminExpensesDetails']:
                if item['name'] in ["audit fee", "auditors' remuneration"]:
                    audit_fee_previous = format_number(item['value'], is_cost_or_admin=False)
                    break

            d_salary_curr = "-"
            d_salary_prev = "-"
            for item in statement_current['GeneralAdminExpensesDetails']:
                if item['name'] in ["director's remuneration", "director’s remuneration"]:
                    d_salary_curr = format_number(item['value'], is_cost_or_admin=False)
                    break
            for item in statement_previous['GeneralAdminExpensesDetails']:
                if item['name'] in ["director's remuneration", "director’s remuneration"]:
                    d_salary_prev = format_number(item['value'], is_cost_or_admin=False)
                    break

            benefit_current = 0
            benefit_previous = 0
            for item in statement_current['GeneralAdminExpensesDetails']:
                if item['name'] in ["director's remuneration", "director’s remuneration"]:
                    benefit_current += item['value']
                elif item['name'] == "salaries":
                    benefit_current += item['value']
            for item in statement_previous['GeneralAdminExpensesDetails']:
                if item['name'] in ["director's remuneration", "director’s remuneration"]:
                    benefit_previous += item['value']
                elif item['name'] == "salaries":
                    benefit_previous += item['value']
            benefit_current = format_number(benefit_current, is_cost_or_admin=False) if benefit_current != 0 else "-"
            benefit_previous = format_number(benefit_previous, is_cost_or_admin=False) if benefit_previous != 0 else "-"

            due_from_director_curr = 0
            due_to_director_curr = 0
            due_from_director_prev = 0
            due_to_director_prev = 0
            show_due_paragraph = False


            inventories_curr = 0
            inventories_prev = 0
            for item in balance_current['current_assets']:
                if item['name'] == "inventories":
                    inventories_curr = item['value']
                    break
            for item in balance_previous['current_assets']:
                if item['name'] == "inventories":
                    inventories_prev = item['value']
                    break
            inventories_curr = format_number(inventories_curr, is_cost_or_admin=False, is_liability=False) if inventories_curr != 0 else "-"
            inventories_prev = format_number(inventories_prev, is_cost_or_admin=False, is_liability=False) if inventories_prev != 0 else "-"

            investment_in_sub_curr = 0
            investment_in_sub_prev = 0
            subsidiary_items = [
                'investments in a subsidiary', 'investments in subsidiaries',
                'interests in subsidiaries', 'interests in a subsidiary'
            ]
            has_subsidiary_for_report = any(item in all_items for item in subsidiary_items)
            if has_subsidiary_for_report:
                for item in balance_current['non_current_assets']:
                    if item['name'] in subsidiary_items:
                        investment_in_sub_curr += item['value']
                for item in balance_previous['non_current_assets']:
                    if item['name'] in subsidiary_items:
                        investment_in_sub_prev += item['value']
            investment_in_sub_curr = format_number(investment_in_sub_curr, is_cost_or_admin=False, is_liability=False) if investment_in_sub_curr != 0 else "-"
            investment_in_sub_prev = format_number(investment_in_sub_prev, is_cost_or_admin=False, is_liability=False) if investment_in_sub_prev != 0 else "-"

            investment_in_asso_curr = 0
            investment_in_asso_prev = 0
            associate_items = ['investments in an associate', 'investments in associates']
            for item in balance_current['non_current_assets']:
                if item['name'] in associate_items:
                    investment_in_asso_curr += item['value']
            for item in balance_previous['non_current_assets']:
                if item['name'] in associate_items:
                    investment_in_asso_prev += item['value']
            investment_in_asso_curr = format_number(investment_in_asso_curr, is_cost_or_admin=False, is_liability=False) if investment_in_asso_curr != 0 else "-"
            investment_in_asso_prev = format_number(investment_in_asso_prev, is_cost_or_admin=False, is_liability=False) if investment_in_asso_prev != 0 else "-"

            dividend_curr = 0
            dividend_prev = 0
            dividend_items = ['dividends paid to shareholders', 'dividends paid to a shareholder', 'dividends paid to shareholder']
            for item in balance_current['equity']:
                if item['name'] in dividend_items:
                    dividend_curr += item['value']
            for item in balance_previous['equity']:
                if item['name'] in dividend_items:
                    dividend_prev += item['value']

            shares_gap_formatted = format_number(int(shares_curr) - int(shares_prev), is_cost_or_admin=False, is_liability=False)
            shares_curr_formatted = format_number(shares_curr, is_cost_or_admin=False, is_liability=False)
            shares_prev_formatted = format_number(shares_prev, is_cost_or_admin=False, is_liability=False)

            shares_cap_curr = 0
            shares_cap_prev = 0
            for item in balance_current['equity']:
                if item['name'] == "share capital":
                    shares_cap_curr = item['value']
                    break
            for item in balance_previous['equity']:
                if item['name'] == "share capital":
                    shares_cap_prev = item['value']
                    break
            shares_cap_gap = int(shares_cap_curr) - int(shares_cap_prev)

            for item in balance_current['current_assets']:
                if item['name'] in ["amount due from a director", "amount due from the director", "amount due from director", "amount due from directors"]:
                    due_from_director_curr = item['value']
                    show_due_paragraph = True
                    break
            for item in balance_current['current_liabilities']:
                if item['name'] in ["amount due to a director", "amount due to the director", "amount due to director", "amount due to directors"]:
                    due_to_director_curr = -item['value']
                    break
            for item in balance_previous['current_assets']:
                if item['name'] in ["amount due from a director", "amount due from the director", "amount due from director", "amount due from directors"]:
                    due_from_director_prev = item['value']
                    show_due_paragraph = True
                    break
            for item in balance_previous['current_liabilities']:
                if item['name'] in ["amount due to a director", "amount due to the director", "amount due to director", "amount due to directors"]:
                    due_to_director_prev = -item['value']
                    break

            if due_from_director_curr != 0:
                due_curr = due_from_director_curr
            else:
                due_curr = due_to_director_curr

            if due_from_director_prev != 0:
                due_prev = due_from_director_prev
            else:
                due_prev = due_to_director_prev

            if self._first_year:
                due_max = max(due_curr, shares_cap_curr)
            else:
                due_max = max(due_curr, due_prev)

            due_curr = format_number(due_curr, is_cost_or_admin=False, is_liability=False) if due_curr != 0 else "-"
            due_prev = format_number(due_prev, is_cost_or_admin=False, is_liability=False) if due_prev != 0 else "-"
            due_max = format_number(due_max, is_cost_or_admin=False, is_liability=False) if due_max != 0 else "-"

            cap_res_curr = 0
            cap_res_prev = 0
            cap_res_gap = 0
            for item in balance_current['equity']:
                if item['name'] in ["capital reserves", "reserves"]:
                    cap_res_curr = item['value']
                    break
            for item in balance_previous['equity']:
                if item['name'] in ["capital reserves", "reserves"]:
                    cap_res_prev = item['value']
                    break
            cap_res_gap = int(cap_res_curr) - int(cap_res_prev)
            cap_res_curr = format_number(cap_res_curr, is_cost_or_admin=False, is_liability=False)
            cap_res_prev = format_number(cap_res_prev, is_cost_or_admin=False, is_liability=False)

            balance_before_current = self._accountant_helper._get_balance_before_period(current_year)
            balance_before_previous = self._accountant_helper._get_balance_before_period(previous_year)
            profit_for_year_current = statement_current['ProfitForYear']
            profit_for_year_previous = statement_previous['ProfitForYear']

            re_curr_num = balance_before_current + profit_for_year_current + dividend_curr 
            re_prev_num = balance_before_previous + profit_for_year_previous + dividend_prev


            if self._first_year:
                if re_curr_num >= 0:
                    re_name = "Retained earnings"
                else:
                    re_name = "Accumulated loss"
                re_name_fn = re_name
            else:
                if re_curr_num > 0 and re_prev_num > 0:
                    re_name = "Retained earnings"
                elif re_curr_num < 0 and re_prev_num < 0:
                    re_name = "Accumulated loss"
                else:
                    re_name = "Retained earnings/(accumulated loss)"
                if re_curr_num >= 0:
                    re_name_fn = "Retained earnings"
                else:
                    re_name_fn = "Accumulated loss"

            re_curr = format_number(re_curr_num, is_cost_or_admin=False, is_liability=False)
            re_prev = format_number(re_prev_num, is_cost_or_admin=False, is_liability=False)

            total_equity_current = balance_current['total_equity'] + re_curr_num - dividend_curr
            total_equity_previous = balance_previous['total_equity'] + re_prev_num - dividend_prev
            profit_current = abs(profit_for_year_current)
            equity_current = abs(total_equity_current)

            re_total = re_prev_num + profit_for_year_current + dividend_curr
            re_total2 = total_equity_previous + cap_res_gap + profit_for_year_current + dividend_prev

            if net_assets_current != total_equity_current:
                logger.error(f"NetAssetsCurrent ({net_assets_current}) does not equal TotalEquityCurrent ({total_equity_current})")
                raise NetAssetsEquityMismatchError(
                    f"NetAssetsCurrent ({net_assets_current}) does not equal TotalEquityCurrent ({total_equity_current}). Document generation aborted."
                )

            due_final_holding_parent_company_info = self.get_due_info(self.due_from_final_holding_parent_company_items,
                                                                      self.due_to_final_holding_parent_company_items,
                                                                      self.due_final_holding_parent_company_items)
            due_shareholder_info = self.get_due_info(self.due_from_shareholder_items, self.due_to_shareholder_items, self.due_shareholder_items)
            due_imme_parent_company_info = self.get_due_info(self.due_from_imme_parent_company_items, self.due_to_imme_parent_company_items, self.due_imme_parent_company_items)
            due_ultimate_holding_company_info = self.get_due_info(self.due_from_ultimate_holding_company_items, self.due_to_ultimate_holding_company_items, self.due_ultimate_holding_company_items)
            due_holding_company_info = self.get_due_info(self.due_from_holding_company_items, self.due_to_holding_company_items, self.due_holding_company_items)
            footnote_vars = [
                ('HasProperty', has_property, ['property, plant and equipment']),
                ('HasLongTermInvestments', has_long_term_investments, ['long-term investments']),
                ('HasCurrentInvestments', has_current_investments, ['current investments']),
                ('HasInventories', has_inventories, ['inventories']),
                ('HasIntangibleAsset', has_intangible_asset, ['intangible assets']),
                ('HasSubsidiary', has_subsidiary_for_report, ['investments in a subsidiary', 'investments in subsidiaries', 'interests in subsidiaries', 'interests in a subsidiary']),
                ('HasAssociate', has_associate, ['investments in an associate', 'investments in associates']),
                ('HasDueFromFinalParent', due_final_holding_parent_company_info['need_footnote'],
                 self.due_final_holding_parent_company_items),
                ('HasDueFromHoldingCompany', due_holding_company_info['need_footnote'], self.due_holding_company_items),
                ('HasDueFromImmeParent', due_imme_parent_company_info['need_footnote'], self.due_imme_parent_company_items),
                ('HasUltimateCompany', due_ultimate_holding_company_info['need_footnote'], self.due_ultimate_holding_company_items),
                ('HasDueFromShareHolder', due_shareholder_info['need_footnote'], 
                self.due_shareholder_items),
                ('HasRelatedParty', has_related_party, []),
                ('HasUltimateCompany', has_ultimate_company, [])
            ]
            
            logger.debug("Assigning footnote numbers...")
            footnote_numbers = {}
            current_footnote = 10
            if 'share capital' in all_items:
                footnote_numbers['share capital'] = '8,9'
                logger.debug("Assigned Share capital: 8,9")

            director_items = [
                'amount due from a director', 
                'amount due from the director', 
                'amount due from director', 
                'amount due from directors',
                'amount due to a director', 
                'amount due to the director', 
                'amount due to director', 
                'amount due to directors'
            ]
            has_director_items = any(item in all_items for item in director_items)
            need_footnote = 'amount due from a director' in all_items or 'amount due from directors' in all_items or 'amount due from director' in all_items or 'amount due from the director' in all_items
            if need_footnote and has_director_items:
                for item in director_items:
                    if item in all_items:
                        footnote_numbers[item] = 7
                        logger.debug(f"Assigned {item}: 7")

            for var_name, var_value, item_names in footnote_vars:
                logger.debug(f"Checking {var_name}: {var_value}")
                if var_value:
                    if item_names:
                        for item_name in item_names:
                            footnote_numbers[item_name] = current_footnote
                            logger.debug(f"Assigned {item_name}: {current_footnote}")
                    else:
                        footnote_numbers[var_name] = current_footnote
                        logger.debug(f"Assigned {var_name}: {current_footnote}")
                    current_footnote += 1
            logger.debug(f"Final footnote_numbers: {footnote_numbers}")

            non_current_asset_names = sorted(set(
                item['name'] for item in balance_current['non_current_assets']
            ).union(
                item['name'] for item in balance_previous['non_current_assets']
            ))

            non_current_assets_list = []
            for idx, name in enumerate(non_current_asset_names):
                current_value = next((item['value'] for item in balance_current['non_current_assets'] if item['name'] == name), 0)
                prev_value = next((item['value'] for item in balance_previous['non_current_assets'] if item['name'] == name), 0)
                fnnum = str(footnote_numbers.get(name, ""))
                logger.debug(f"Item: {name}, Footnote: '{fnnum}'")
                non_current_assets_list.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value),
                    'pr': format_number(prev_value),
                    'fnnum': fnnum,
                    'is_last': idx == len(non_current_asset_names) - 1
                })

            current_asset_names = sorted(set(
                item['name'] for item in balance_current['current_assets']
            ).union(
                item['name'] for item in balance_previous['current_assets']
            ))

            current_assets_list = []
            for idx, name in enumerate(current_asset_names):
                current_value = next((item['value'] for item in balance_current['current_assets'] if item['name'] == name), 0)
                prev_value = next((item['value'] for item in balance_previous['current_assets'] if item['name'] == name), 0)
                fnnum = str(footnote_numbers.get(name, ""))
                logger.debug(f"Item: {name}, Footnote: '{fnnum}'")
                current_assets_list.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value),
                    'pr': format_number(prev_value),
                    'fnnum': fnnum,
                    'is_last': idx == len(current_asset_names) - 1
                })

            current_liabilities_names = sorted(set(
                item['name'] for item in balance_current['current_liabilities']
            ).union(
                item['name'] for item in balance_previous['current_liabilities']
            ))

            current_liabilities_list = []
            for idx, name in enumerate(current_liabilities_names):
                current_value = next((item['value'] for item in balance_current['current_liabilities'] if item['name'] == name), 0)
                prev_value = next((item['value'] for item in balance_previous['current_liabilities'] if item['name'] == name), 0)
                fnnum = str(footnote_numbers.get(name, ""))
                logger.debug(f"Item: {name}, Footnote: '{fnnum}'")
                current_liabilities_list.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value, is_liability=True),
                    'pr': format_number(prev_value, is_liability=True),
                    'fnnum': fnnum,
                    'is_last': idx == len(current_liabilities_names) - 1
                })

            non_current_liabilities_names = sorted(set(
                item['name'] for item in balance_current['non_current_liabilities']
            ).union(
                item['name'] for item in balance_previous['non_current_liabilities']
            ))

            non_current_liabilities_list = []
            for idx, name in enumerate(non_current_liabilities_names):
                current_value = next((item['value'] for item in balance_current['non_current_liabilities'] if item['name'] == name), 0)
                prev_value = next((item['value'] for item in balance_previous['non_current_liabilities'] if item['name'] == name), 0)
                fnnum = str(footnote_numbers.get(name, ""))
                logger.debug(f"Item: {name}, Footnote: '{fnnum}'")
                non_current_liabilities_list.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value, is_liability=True),
                    'pr': format_number(prev_value, is_liability=True),
                    'fnnum': fnnum,
                    'is_last': idx == len(non_current_liabilities_names) - 1
                })

            equity_names = set(
                item['name'] for item in balance_current['equity']
            ).union(
                item['name'] for item in balance_previous['equity']
            )

            priority_order = ["share capital", "reserves", "capital reserves"]
            sorted_equity_names = []
            for item in priority_order:
                if item in equity_names:
                    sorted_equity_names.append(item)
                    equity_names.remove(item)
            # DO NOT list dividends
            #sorted_equity_names.extend(sorted(equity_names))

            equity_list = []
            for idx, name in enumerate(sorted_equity_names):
                current_value = next((item['value'] for item in balance_current['equity'] if item['name'] == name), 0)
                prev_value = next((item['value'] for item in balance_previous['equity'] if item['name'] == name), 0)
                fnnum = str(footnote_numbers.get(name, ""))
                logger.debug(f"Item: {name}, Footnote: '{fnnum}'")
                equity_list.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value),
                    'pr': format_number(prev_value),
                    'fnnum': fnnum,
                    'is_last': idx == len(sorted_equity_names) - 1
                })

            self._closing_inventories_curr = 0 
            self._closing_inventories_prev = 0 
            cost_items = []
            cost_items_current = statement_current['CostItemsDetails']
            cost_items_previous = statement_previous['CostItemsDetails']
            cost_item_names = sorted(set(item['name'] for item in cost_items_current + cost_items_previous))
            priority_order = ['opening inventories', 'purchases', 'closing inventories', 'direct costs']
            sorted_cost_item_names = []
            for item in priority_order:
                if item in cost_item_names:
                    sorted_cost_item_names.append(item)
                    cost_item_names.remove(item)
            sorted_cost_item_names.extend(sorted(cost_item_names))
            for idx, name in enumerate(sorted_cost_item_names):
                current_value = next((item['value'] for item in cost_items_current if item['name'] == name), 0)
                previous_value = next((item['value'] for item in cost_items_previous if item['name'] == name), 0)
                if name == 'closing inventories':
                    self._closing_inventories_curr = current_value
                    self._closing_inventories_prev = previous_value
                    cost_items.append({
                        'name': name.capitalize(),
                        'cu': format_number(current_value, is_tax=True),
                        'pr': format_number(previous_value, is_tax=True),
                        'is_last': idx == len(sorted_cost_item_names) - 1
                    })
                else:
                    cost_items.append({
                        'name': name.capitalize(),
                        'cu': format_number(current_value, is_cost_or_admin=True),
                        'pr': format_number(previous_value, is_cost_or_admin=True),
                        'is_last': idx == len(sorted_cost_item_names) - 1
                    })

            turnover_items = []
            turnover_current = statement_current['RevenueItemsDetails']
            turnover_previous = statement_previous['RevenueItemsDetails']
            turnover_names = sorted(set(item['name'] for item in turnover_current + turnover_previous))
            for idx, name in enumerate(turnover_names):
                current_value = next((item['value'] for item in turnover_current if item['name'] == name), 0)
                previous_value = next((item['value'] for item in turnover_previous if item['name'] == name), 0)
                turnover_items.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value),
                    'pr': format_number(previous_value),
                    'is_last': idx == len(turnover_names) - 1
                })

            other_income_items = []
            other_income_current = statement_current['OtherIncomeDetails']
            other_income_previous = statement_previous['OtherIncomeDetails']
            other_income_names = sorted(set(item['name'] for item in other_income_current + other_income_previous))
            for idx, name in enumerate(other_income_names):
                current_value = next((item['value'] for item in other_income_current if item['name'] == name), 0)
                previous_value = next((item['value'] for item in other_income_previous if item['name'] == name), 0)
                other_income_items.append({
                    'name': name.capitalize(),
                    'cu': format_number(current_value),
                    'pr': format_number(previous_value),
                    'is_last': idx == len(other_income_names) - 1
                })

            general_admin_expenses_items = []
            general_admin_current = statement_current['GeneralAdminExpensesDetails']
            general_admin_previous = statement_previous['GeneralAdminExpensesDetails']
            general_admin_names = sorted(set(item['name'] for item in general_admin_current + general_admin_previous))
            for idx, name in enumerate(general_admin_names):
                current_value = next((item['value'] for item in general_admin_current if item['name'] == name), 0)
                previous_value = next((item['value'] for item in general_admin_previous if item['name'] == name), 0)
                if current_value > 0:
                    current_value_fmt = format_number(current_value, is_cost_or_admin=True)
                else:
                    current_value_fmt = format_number(current_value, is_tax=True)
                if previous_value > 0:
                    previous_value_fmt = format_number(previous_value, is_cost_or_admin=True)
                else:
                    previous_value_fmt = format_number(previous_value, is_tax=True)
                general_admin_expenses_items.append({
                    'name': name.capitalize(),
                    'cu': current_value_fmt,
                    'pr': previous_value_fmt,
                    'is_last': idx == len(general_admin_names) - 1
                })

            show_gross_profit = statement_current['GrossProfit'] != 0 or statement_previous['GrossProfit'] != 0
            logger.debug(f'cost_items: {cost_items}')
            data = {
                "CompanyNameInChinesePlaceholder": company_name_cn,
                "CompanyNameInEnglishPlaceholder": self._company_name_en,
                "CompanyAddressPlaceHolder": company_address_cleaned,
                "BusinessType": self._business_type,
                "BusinessDescriptionPlaceholder": business_description,
                "bizAdditionalDesc": additional_business_description,
                "LastDayOfYear": last_day_of_year,
                "LastDayOfYearNum": self._last_day_date_num,
                "OneYearAgo": self._one_year_ago,
                "ExactlyOneYearAgo": self._exactly_one_year_ago,
                "DateOfIncorporation": self._date_of_incorporation,
                "AuditFirmInEnglishPlacehoder": audit_firm,
                "ApprovalDatePlaceholder": approval_date,
                "AuditorNamePlaceholder": auditor_name,
                "AuditorLicenseNoPlaceholder": auditor_license,
                "CuYear": self._cu_year_for_first_year if self._first_year else str(current_year),
                "PrYear": self._pr_year_for_first_year if self._first_year else str(previous_year),
                "CYear": self._c_year_for_first_year, 
                "Exceed18m": self._over18m,
                "CurrencyDesc": currency_desc,
                "CurrencyFullDesc": currency_full_desc,
                "Currency": self._currency,
                "directors": self._directors_list,
                "FirstDirectorNamePlaceholder": first_director_name,
                "BusinessType": self._business_type,
                "RevenueName": revenue_name,
                "RevenueCurrent": format_number(statement_current['Revenue']),
                "CostSalesCurr": format_number(statement_current['CostOfSales'], is_cost_or_admin=True),
                "CostSalesPrev": format_number(statement_previous['CostOfSales'], is_cost_or_admin=True),
                "CostOfSalesCurrent": format_number(statement_current['CostOfSales'], is_cost_or_admin=True),
                "GrossProfitCurrent": format_number(statement_current['GrossProfit']),
                "OtherIncomeCurrent": format_number(statement_current['OtherIncome']),
                "GeneralAdminExpensesCurrent": format_number(statement_current['GeneralAdminExpenses'], is_cost_or_admin=True),
                "FinanceCostsCurrent": format_number(statement_current['FinanceCosts']),
                "CalcTotalCurrent": format_number(statement_current['CalcTotal']),
                "ProfitBeforeTaxCurrent": format_number(statement_current['ProfitBeforeTax']),
                "TaxationCurrent": format_number(statement_current['Taxation']),
                "TaxationCurrentFn": format_number(statement_current['Taxation'], is_tax=True),
                "ProfitForYearCurrent": format_number(statement_current['ProfitForYear']),
                "RevenuePrevious": format_number(statement_previous['Revenue']),
                "CostOfSalesPrevious": format_number(statement_previous['CostOfSales'], is_cost_or_admin=True),
                "GrossProfitPrevious": format_number(statement_previous['GrossProfit']),
                "OtherIncomePrevious": format_number(statement_previous['OtherIncome']),
                "GeneralAdminExpensesPrevious": format_number(statement_previous['GeneralAdminExpenses'], is_cost_or_admin=True),
                "FinanceCostsPrevious": format_number(statement_previous['FinanceCosts']),
                "CalcTotalPrevious": format_number(statement_previous['CalcTotal']),
                "ProfitBeforeTaxPrevious": format_number(statement_previous['ProfitBeforeTax']),
                "TaxationPrevious": format_number(statement_previous['Taxation']),
                "TaxationPreviousFn": format_number(statement_previous['Taxation'], is_tax=True),
                "ProfitForYearPrevious": format_number(statement_previous['ProfitForYear']),
                "PLBFTName": plbft_name,
                "PLName": pl_name,
                "PLNameCurr": pl_name_curr,
                "PLNameCap": pl_name_cap,
                "GrossPLName": gross_pl_name,
                "non_current_assets": non_current_assets_list,
                "current_assets": current_assets_list,
                "current_liabilities": current_liabilities_list,
                "non_current_liabilities": non_current_liabilities_list,
                "equity": equity_list,
                "TotalNonCurrentAssetsCurrent": format_number(balance_current['total_non_current_assets']),
                "TotalNonCurrentAssetsPrevious": format_number(balance_previous['total_non_current_assets']),
                "TotalCurrentAssetsCurrent": format_number(balance_current['total_current_assets']),
                "TotalCurrentAssetsPrevious": format_number(balance_previous['total_current_assets']),
                "TotalCurrentLiabilitiesCurrent": format_number(balance_current['total_current_liabilities'], is_liability=True),
                "TotalCurrentLiabilitiesPrevious": format_number(balance_previous['total_current_liabilities'], is_liability=True),
                "TotalNonCurrentLiabilitiesCurrent": format_number(balance_current['total_non_current_liabilities'], is_liability=True),
                "TotalNonCurrentLiabilitiesPrevious": format_number(balance_previous['total_non_current_liabilities'], is_liability=True),
                "NetAssetsCurrent": format_number(balance_current['net_assets']),
                "NetAssetsPrevious": format_number(balance_previous['net_assets']),
                "TotalEquityCurrent": format_number(total_equity_current), #format_number(balance_current['total_equity'] + re_curr_num),
                "TotalEquityPrevious": format_number(total_equity_previous), #format_number(balance_previous['total_equity'] + re_prev_num),
                "ProfitCurrent": format_number(profit_current, is_cost_or_admin=False, is_liability=False),
                "EquityCurrent": format_number(equity_current, is_cost_or_admin=False, is_liability=False),
                "NetAssetsName": net_assets_name,
                "turnover_items": turnover_items,
                "cost_items": cost_items,
                "show_cost_of_sales": len(cost_items) > 0,
                "show_gross_profit": show_gross_profit,
                "show_non_current_assets": len(non_current_assets_list) > 0,
                "show_non_current_liabilities": len(non_current_liabilities_list) > 0,
                "other_income_items": other_income_items,
                "show_other_income": len(other_income_items) > 0,
                "general_admin_expenses_items": general_admin_expenses_items,
                "show_negative_net_assets": net_assets_current < 0,
                "cash_bank": cash_bank,
                "AuditFeeCurrent": audit_fee_current,
                "Shareholders": shareholders,
                "AuditFeePrevious": audit_fee_previous,
                "DSalaryCurr": d_salary_curr,
                "DSalaryPrev": d_salary_prev,
                "BenefitCurrent": benefit_current,
                "BenefitPrevious": benefit_previous,
                "show_due_paragraph": show_due_paragraph,
                "DuePrev": due_prev,
                "DueCurr": due_curr,
                "DueMax": due_max,
                "SharesCurr": shares_curr_formatted,
                "SharesPrev": shares_prev_formatted,
                "SharesGap": shares_gap_formatted,
                "SharesCapCurr": format_number(shares_cap_curr),
                "SharesCapPrev": format_number(shares_cap_prev),
                "SharesCapGap": format_number(shares_cap_gap),
                "CapResCurr": cap_res_curr,
                "CapResPrev": cap_res_prev,
                "CapResGap": format_number(cap_res_gap),
                "RECurr": re_curr,
                "REPrev": re_prev,
                "RETotal": format_number(re_total),
                "RETotal2": format_number(re_total2),
                "REName": re_name,
                "RENameFn": re_name_fn,
                "HasNameChanged": has_name_changed,
                "PassedDate": passed_date,
                "NewCompanyName": new_company_name,
                "EffectiveDate": effective_date,
                "OldCompanyName": old_company_name,
                "HasRelatedParty": has_related_party,
                "InventoryValuation": inventory_valuation,
                "TaxOpt": tax_opt,
                "CapitalIncrease": capital_increase,
                "HasUltimateCompany": has_ultimate_company,
                "UltimateCompanyOption": ultimate_company_option,
                "UltimateCompanyName1": ultimate_company_name1,
                "UltimateCompanyLocation1": ultimate_company_location1,
                "UltimateCompanyName2": ultimate_company_name2,
                "UltimateCompanyLocation2": ultimate_company_location2,
                "DueFromDir": due_from_dir,
                "HasSubsidiary": has_subsidiary_for_report,
                "HasServiceFeeIncome": has_service_fee_income,
                "HasAgencyFeeIncome": has_agency_fee_income,
                "HasSalesOfGoods": has_sales_of_goods,
                "HasLongTermInvestments": has_long_term_investments,
                "HasProperty": has_property,
                "HasInvestment": has_investment,
                "HasInventories": has_inventories,
                "HasIntangibleAsset": has_intangible_asset,
                "HasReserve": has_reserve,
                "HasCurrentInvestments": has_current_investments,
                "HasIntangibleAssets": has_intangible_assets,
                "HasAssociate": has_associate,
                "HasDueFromFinalParent": due_final_holding_parent_company_info['need_footnote'],
                "HasDueToFinalParent2": due_final_holding_parent_company_info['both_to'], #self.exists_due_to_final_parent_for_years(),
                "HasDueFromImmeParent": due_imme_parent_company_info['need_footnote'],
                "HasDueToImmeParent2": due_imme_parent_company_info['both_to'],
                "HasDueFromShareHolder": due_shareholder_info['need_footnote'],
                "HasDueToShareHolder2": due_shareholder_info['both_to'], #self.exists_due_to_shareholder(),
                "DueFromShareHolderName": due_shareholder_info['title_name'], #due_from_shareholder_name,
                "SubsidiaryName": "" if not has_subsidiary_for_report else (
                    next((item for item in subsidiary_items if item.lower() in all_items), "")
                    if len([item for item in subsidiary_items if item.lower() in all_items]) == 1
                    else ""
                ),
                "InventoriesCurr": inventories_curr,
                "InventoriesPrev": inventories_prev,
                "InvestmentInSubCurr": investment_in_sub_curr,
                "InvestmentInSubPrev": investment_in_sub_prev,
                "InvestmentInAssoCurr": investment_in_asso_curr,
                "InvestmentInAssoPrev": investment_in_asso_prev,
                "DueFinalParentName": due_final_holding_parent_company_info['title_name'], #due_final_parent_name,
                "DueFinalCurr": due_final_holding_parent_company_info['curr'], #due_final_curr,
                "DueFinalPrev": due_final_holding_parent_company_info['prev'], #due_final_curr,due_final_prev,
                "DueFinalMax": due_final_holding_parent_company_info['max'], #due_final_curr,due_final_prev,due_final_max,
                "DueImmeParentName": due_imme_parent_company_info['title_name'], 
                "DueImmeCurr": due_imme_parent_company_info['curr'],
                "DueImmePrev": due_imme_parent_company_info['prev'],
                "DueImmeMax": due_imme_parent_company_info['max'],
                "DueShareholderCurr": due_shareholder_info['curr'],
                "DueShareholderPrev": due_shareholder_info['prev'], #due_shareholder_prev,
                "DueShareholderMax": due_shareholder_info['max'], #due_shareholder_max,
                "DividendCurr": format_number(dividend_curr, is_cost_or_admin=False, is_liability=False) if dividend_curr != 0 else "-",
                "DividendPrev": format_number(dividend_prev, is_cost_or_admin=False, is_liability=False) if dividend_prev != 0 else "-",
                "InvestmentInCompany": investment_in_company,
                "InvestmentInSecurity": investment_in_security,
                "Opinion": audit_opinion,
                "LongTermInvestmentCurr": format_number(long_term_investments_curr),
                "LongTermInvestmentPrev": format_number(long_term_investments_prev),
                "CurrentInvestmentCurr": format_number(current_investment_curr),
                "CurrentInvestmentPrev": format_number(current_investment_prev),
                "HasDueFromUltimateHoldingCompany": due_ultimate_holding_company_info['need_footnote'],
                "HasDueToUltimateHoldingCompany2": due_ultimate_holding_company_info['both_to'],
                "DueFromUltimateHoldingCompanyName": due_ultimate_holding_company_info['title_name'],
                "HasDueFromHolding": due_holding_company_info['need_footnote'],
                "DueFromHoldingName": due_holding_company_info['title_name'],
                "HasDueToHolding2": due_holding_company_info['both_to'],
                "AuditType": audit_type,
                "DueUltimateHoldingCompanyCurr": due_ultimate_holding_company_info['curr'],
                "DueUltimateHoldingCompanyPrev": due_ultimate_holding_company_info['prev'],
                "DueUltimateHoldingCompanyMax": due_ultimate_holding_company_info['max'],
                "DueHoldingCompanyCurr": due_holding_company_info['curr'],
                "DueHoldingCompanyPrev": due_holding_company_info['prev'],
                "DueHoldingCompanyMax": due_holding_company_info['max'],
            }

            logger.debug(f"SubsidiaryName: {data['SubsidiaryName']}")
            excluded_fields = [
                "DueFromShareHolderName", "CompanyNameInChinesePlaceholder", "bizAdditionalDesc",
                "DueFinalParentName", "DueImmeParentName", "DueFromShareHolderName", "SubsidiaryName"
            ]
            if not self._first_year:
                excluded_fields.extend(["DateOfIncorporation", "CYear"])
            if not has_name_changed:
                excluded_fields.extend(["PassedDate", "NewCompanyName", "OldCompanyName", "EffectiveDate"])
            if has_ultimate_company:
                if ultimate_company_option == "option1":
                    excluded_fields.extend(["UltimateCompanyName2", "UltimateCompanyLocation2"])
                elif ultimate_company_option == "option3":
                    excluded_fields.extend(["UltimateCompanyLocation2"])
            else:
                excluded_fields.extend(["UltimateCompanyOption", "UltimateCompanyName1", "UltimateCompanyLocation1", "UltimateCompanyName2", "UltimateCompanyLocation2"])
            excluded_fields.extend(["DueFromUltimateHoldingCompanyName"])
            excluded_fields.extend(["DueFromHoldingName"])

            for key, value in data.items():
                if key not in excluded_fields and isinstance(value, str) and not value.strip():
                    logger.error(f"Please fill in all fields: {key}")
                    return None, f"Error: Please fill in all fields: {key}"
            if not self._directors_list:
                logger.error("Please provide at least one director name")
                return None, "Error: Please provide at least one director name."
            if has_long_term_investments:
                if not investment_in_company and not investment_in_security:
                    logger.error("Please select investment in company or security")
                    return None, "Error: Please select investment in company or security."

            final_output_path = output_path if output_path else "audit_report_filled.docx"
            logger.info(f"Will save output to: {final_output_path}")

            logger.info("Rendering template")
            template.render(data)
            logger.info("Rendering template completed")
            template.save(final_output_path)

            doc = Document(final_output_path)
            #insert_page_break_before_income_statement(doc)

            for table in doc.tables:
                for row in table.rows:
                    row.height = None
                    for cell in row.cells:
                        cell_text = "".join(p.text for p in cell.paragraphs).strip()
                        has_underline = '[[UNDERLINE]]' in cell_text
                        has_dbline = '[[DBLine]]' in cell_text

                        # Log cell text for debugging
                        logger.debug(f"Processing cell with text: '{cell_text}', has_dbline: {has_dbline}, has_underline: {has_underline}")

                        # Remove [[DBLine]] and [[UNDERLINE]] from all paragraphs and runs
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if '[[DBLine]]' in run.text or '[[UNDERLINE]]' in run.text:
                                    run.font.size = Pt(11)
                                    run.text = run.text.replace('[[DBLine]]', '').replace('[[UNDERLINE]]', '')
                                    logger.debug(f"Removed placeholders from run: '{run.text}'")

                            # Re-check paragraph text after run replacements
                            paragraph_text = paragraph.text
                            if '[[DBLine]]' in paragraph_text or '[[UNDERLINE]]' in paragraph_text:
                                #run.font.size = Pt(11)
                                clean_text = paragraph_text.replace('[[DBLine]]', '').replace('[[UNDERLINE]]', '')
                                #clean_text.font.size = Pt(11)
                                # Clear all runs
                                while paragraph.runs:
                                    paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
                                # Add cleaned text as a new run
                                paragraph.add_run(clean_text)
                                logger.debug(f"Replaced paragraph text with: '{clean_text}'")

                            paragraph.space_before = 0
                            paragraph.space_after = 0
                            paragraph.style.paragraph_format.line_spacing = 1.0

                        # Apply borders if needed
                        if has_dbline or has_underline:
                            tc = cell._element
                            tcPr = tc.get_or_add_tcPr()
                            # Remove existing borders
                            existing_borders = tcPr.xpath('./w:tcBorders')
                            if existing_borders:
                                tcPr.remove(existing_borders[0])
                            # Add new borders
                            tcBorders = OxmlElement('w:tcBorders')
                            bottom = OxmlElement('w:bottom')
                            if has_dbline:
                                bottom.set(qn('w:val'), 'double')
                                bottom.set(qn('w:sz'), '8')  # Size in eighths of a point
                                logger.debug("Applying double bottom border")
                            elif has_underline:
                                bottom.set(qn('w:val'), 'single')
                                bottom.set(qn('w:sz'), '4')
                                logger.debug("Applying single bottom border")
                            bottom.set(qn('w:space'), '0')
                            bottom.set(qn('w:color'), '000000')
                            tcBorders.append(bottom)
                            tcPr.append(tcBorders)

            logger.info("Before update_fields")
            update_fields(doc)
            logger.info("Before final save")
            doc.save(final_output_path)
            logger.info("Document saved successfully")
            if self._inventories_curr != self._closing_inventories_curr or self._inventories_prev != self._closing_inventories_prev:
                warning = f"inventories mismatch:\n, inventories_curr: {self._inventories_curr}, inventories_prev: {self._inventories_prev}\n closing_inventories_curr: {self._closing_inventories_curr}, closing_inventories_prev: {self._closing_inventories_prev}"
                return True, warning
            return True, ""
        except InvalidTBSheetFormatError as e:
            logger.error(f"Invalid trial balance sheet format: {str(e)}")
            return None, str(e)
        except UnrecognizedItemError as e:
            logger.error(f"Unrecognized item: {str(e)}")
            import traceback
            traceback.print_exc()
            return None, str(e)
        except InvalidItemNameError as e:
            logger.error(f"Invalid item name: {str(e)}")
            return None, str(e)
        except NetAssetsEquityMismatchError as e:
            logger.error(f"Net assets equity mismatch: {str(e)}")
            return None, str(e)
        except FileNotFoundError as e:
            logger.error(f"File not found: {str(e)}")
            return None, str(e)
        except ValueError as e:
            logger.error(f"Value error: {str(e)}")
            import traceback
            traceback.print_exc()
            return None, str(e)
        except TypeError as e:
            logger.exception(f"Type error during document generation: {str(e)}")
            return None, str(e)
        except Exception as e:
            logger.exception(f"Unexpected error during document generation: {str(e)}")
            return None, str(e)